#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar la documentación del proyecto FitMeal en formato Word.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

LOGO_PATH = os.path.join(os.path.dirname(__file__), "frontend", "public", "FitMeal_logoblanco.png")
LOGO_COLOR_PATH = os.path.join(os.path.dirname(__file__), "frontend", "public", "FitMeal_logoblanco.png")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "FitMeal_Documentacion.docx")

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def set_heading_style(paragraph, level=1):
    """Apply heading formatting."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x3460, 0x01)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x53, 0x35, 0x83)
            run.font.size = Pt(12)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_screenshot_placeholder(doc, caption="[CAPTURA DE PANTALLA]", width_cm=14):
    """Add a bordered box to indicate where a screenshot goes."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'A0A0A0')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"\n\n{caption}\n\n")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()  # spacing after

def add_code_block(doc, code_text):
    """Add a styled code block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    tcPr.append(shd)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    para = cell.paragraphs[0]
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    doc.add_paragraph()

def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_info_table(doc, data, headers=None):
    """Add a formatted table with optional headers."""
    cols = len(data[0])
    table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_offset = 0
    if headers:
        row_offset = 1
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1A1A2E')
            tcPr.append(shd)
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.cell(i + row_offset, j)
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                tcPr.append(shd)
    doc.add_paragraph()

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

# ─────────────────────────────────────────────
# Header with logo on every page
# ─────────────────────────────────────────────

def add_header_with_logo(doc):
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(1, 2, Inches(6))
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Left cell: logo
    left_cell = header_table.cell(0, 0)
    left_cell.width = Inches(1.5)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        run = left_para.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.4))
    else:
        run = left_para.add_run("FitMeal")
        run.font.bold = True
        run.font.size = Pt(14)
    # Right cell: title
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_para.add_run("FitMeal – Documentació Tècnica del Projecte")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─────────────────────────────────────────────
# DOCUMENT BUILD
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_header_with_logo(doc)

    # ─── 0. COVER PAGE ───────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo centered
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(3.5))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Gestió de Nutrició i Entrenament Personalitzat Basat en Web")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FitMeal – La teva app de fitness i nutrició")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x53, 0x35, 0x83)

    doc.add_paragraph()
    doc.add_paragraph()

    info = [
        ("Autors:", "Kevin Castellón Ledezma  ·  Tommy  ·  [Nom 3]"),
        ("Director:", "[Nom del director del projecte]"),
        ("Curs:", "2n DAW – 2025 / 2026"),
        ("Centre:", "Centre d'Estudis Monlau"),
        ("Data:", "Maig 2026"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.font.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)

    doc.add_page_break()

    # ─── 1. DEDICATÒRIES ─────────────────────────────────────────────
    add_heading(doc, "1. Dedicatòries i Agraïments", level=1)
    add_body(doc,
        "Volem dedicar aquest projecte a totes les persones que han estat al nostre costat durant "
        "aquests dos anys de formació. Als nostres professors, que han compartit amb nosaltres no "
        "tan sols coneixements tècnics sinó també la passió per la programació. Als companys de curs, "
        "per les hores de col·laboració, els debats i les solucions trobades en comú."
    )
    add_body(doc,
        "Agraïm especialment al nostre director de projecte per la seva orientació constant i per "
        "ajudar-nos a enfocar la complexitat tècnica del projecte. A les nostres famílies, per la "
        "paciència i el suport incondicional durant les nits de codi i els caps de setmana dedicats "
        "al desenvolupament."
    )
    add_body(doc,
        "Finalment, a la comunitat de codi obert i a totes les persones que han contribuït a les "
        "eines que hem fet servir: React, Node.js, MySQL, Tailwind CSS i tantes d'altres sense les "
        "quals aquest projecte no hauria estat possible."
    )
    add_separator(doc)
    doc.add_paragraph()

    # ─── 2. PROPÒSIT ─────────────────────────────────────────────────
    add_heading(doc, "2. Propòsit del Projecte", level=1)
    add_body(doc,
        "FitMeal neix de la necessitat cada cop més present de disposar d'una eina digital que "
        "integri, de manera cohesionada, la gestió de la nutrició i l'entrenament físic personal. "
        "En l'actualitat, la majoria d'aplicacions o bé se centren en la dieta o bé en l'exercici, "
        "però rarament ofereixen les dues dimensions en un únic entorn."
    )
    add_body(doc,
        "El projecte ens atrau per diverses raons. En primer lloc, és un domini real i amb impacte: "
        "qualsevol persona que vulgui millorar la seva salut necessita un pla nutricional i un "
        "entrenament adaptat als seus objectius. En segon lloc, tècnicament és un repte complet: "
        "requereix un back-end robust amb autenticació, gestió de rols, API RESTful i base de dades "
        "relacional, a més d'un front-end modern amb una experiència d'usuari cuidada."
    )
    add_body(doc,
        "A nivell personal, el projecte ens ha permès aplicar tots els coneixements adquirits durant "
        "el cicle en un context pràctic: des del disseny de la base de dades fins al desplegament en "
        "contenidors Docker, passant per la implementació de OAuth, el sistema de rols i la gestió "
        "d'uploads d'imatges. La possibilitat d'oferir una versió premium amb un entrenador personal "
        "virtual ha afegit una capa de complexitat que ha enriquit molt el projecte."
    )
    add_separator(doc)
    doc.add_paragraph()

    # ─── 3. ÍNDEX ─────────────────────────────────────────────────────
    add_heading(doc, "3. Índex", level=1)
    toc_items = [
        ("1", "Dedicatòries i Agraïments", "3"),
        ("2", "Propòsit del Projecte", "3"),
        ("3", "Índex", "4"),
        ("4", "Definició del Projecte", "5"),
        ("4.1", "Requisits Funcionals i No Funcionals", "5"),
        ("4.2", "Disseny Funcional", "6"),
        ("5", "Estudi de Tecnologies i Eines", "8"),
        ("5.1", "Back-end", "8"),
        ("5.2", "Front-end", "9"),
        ("5.3", "Base de Dades", "10"),
        ("5.4", "Infraestructura i Desplegament", "10"),
        ("6", "Implementació", "11"),
        ("6.1", "Procés 1 – Autenticació i Gestió de Rols", "11"),
        ("6.2", "Procés 2 – Gestió de Recetes i Nutrició", "13"),
        ("6.3", "Procés 3 – Gestió d'Exercicis i Entrenament", "15"),
        ("6.4", "Procés 4 – Perfil d'Usuari i Onboarding", "17"),
        ("6.5", "Procés 5 – Botiga i Carret de la Compra", "19"),
        ("6.6", "Procés 6 – Dashboard d'Administrador", "21"),
        ("6.7", "Procés 7 – Dashboard d'Entrenador", "22"),
        ("6.8", "Procés 8 – Sistema de Contacte", "23"),
        ("7", "Millores i Línies de Futur", "24"),
        ("8", "Conclusions", "25"),
        ("9", "Bibliografia / Webgrafia", "26"),
        ("10", "Annexos", "27"),
    ]
    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (num, title, page) in enumerate(toc_items):
        c0 = toc_table.cell(i, 0)
        c0.text = num
        for run in c0.paragraphs[0].runs:
            run.font.bold = True if '.' not in num else False
            run.font.size = Pt(10)
        c0.width = Cm(1.2)
        c1 = toc_table.cell(i, 1)
        c1.text = title
        for run in c1.paragraphs[0].runs:
            run.font.size = Pt(10)
            if '.' not in num:
                run.font.bold = True
        c2 = toc_table.cell(i, 2)
        c2.text = page
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in c2.paragraphs[0].runs:
            run.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_page_break()

    # ─── 4. DEFINICIÓ DEL PROJECTE ───────────────────────────────────
    add_heading(doc, "4. Definició del Projecte", level=1)
    add_body(doc,
        "FitMeal és una aplicació web fullstack que permet als usuaris gestionar la seva alimentació "
        "i entrenament físic de forma personalitzada. La plataforma combina una API RESTful construïda "
        "amb Node.js/Express i una base de dades MySQL, amb un client web reactiu desenvolupat en "
        "React. El sistema diferencia quatre tipus d'usuaris: administradors, usuaris normals, "
        "usuaris premium i entrenadors personals."
    )
    doc.add_paragraph()

    add_heading(doc, "4.1 Requisits Funcionals", level=2)
    requisits = [
        ("RF-01", "El sistema ha de permetre el registre d'usuaris amb email/contrasenya o via OAuth Google."),
        ("RF-02", "El sistema ha de permetre l'inici de sessió i la gestió de sessions amb JWT."),
        ("RF-03", "L'usuari ha de poder completar un onboarding inicial per personalitzar el seu perfil."),
        ("RF-04", "El sistema ha d'oferir un catàleg de recetes amb informació nutricional filtrable."),
        ("RF-05", "L'usuari ha de poder afegir recetes i exercicis als seus favorits."),
        ("RF-06", "El sistema ha d'oferir un catàleg d'exercicis organitzats per grup muscular i dificultat."),
        ("RF-07", "Els usuaris premium han de poder ser assignats a un entrenador personal."),
        ("RF-08", "L'entrenador ha de poder assignar i gestionar rutines per als seus clients."),
        ("RF-09", "L'usuari ha de poder registrar el seu progrés d'entrenament (sèries, repeticions, pes)."),
        ("RF-10", "El sistema ha de disposar d'una botiga amb productes de fitness i un carret de la compra."),
        ("RF-11", "L'administrador ha de poder gestionar (CRUD) usuaris, recetes i exercicis."),
        ("RF-12", "El sistema ha de permetre l'enviament de missatges de contacte per correu electrònic."),
        ("RF-13", "L'usuari ha de poder actualitzar el seu perfil i pujar una foto de perfil."),
        ("RF-14", "El sistema ha d'implementar plans de subscripció (Bàsic, Avançat, Premium)."),
    ]
    add_info_table(doc, requisits, headers=["Codi", "Descripció"])

    add_heading(doc, "Requisits No Funcionals", level=2)
    rnf = [
        ("RNF-01", "Seguretat", "Contrasenyes encriptades amb bcrypt. Tokens JWT amb expiració de 24h. Cookies httpOnly. Rate limiting en endpoints d'autenticació."),
        ("RNF-02", "Rendiment", "Pool de connexions a la base de dades (màx. 10). Paginació en llistats. Build optimitzat amb Vite."),
        ("RNF-03", "Usabilitat", "Interfície responsiva amb Tailwind CSS. Feedback visual amb toast notifications. Animacions fluides amb Framer Motion."),
        ("RNF-04", "Mantenibilitat", "Arquitectura MVC al back-end. Separació de concerns (context, components, pàgines). Variables d'entorn per a configuració."),
        ("RNF-05", "Portabilitat", "Contenidorització amb Docker i Docker Compose. Compatible amb Linux, Windows i macOS."),
        ("RNF-06", "Escalabilitat", "Disseny modular de l'API. Estructura preparada per afegir nous rols o funcionalitats."),
    ]
    add_info_table(doc, rnf, headers=["Codi", "Categoria", "Descripció"])

    doc.add_paragraph()
    add_heading(doc, "4.2 Disseny Funcional – Arquitectura General", level=2)
    add_body(doc,
        "L'arquitectura del sistema segueix el patró client-servidor. El front-end és una SPA (Single "
        "Page Application) construïda amb React que es comunica amb el back-end a través d'una API "
        "RESTful. El back-end segueix el patró MVC (Model-View-Controller) amb Express, i persisteix "
        "les dades en una base de dades relacional MySQL."
    )
    doc.add_paragraph()
    add_screenshot_placeholder(doc, "[ DIAGRAMA D'ARQUITECTURA: Client (React SPA) ↔ API REST (Node.js/Express) ↔ MySQL ]", width_cm=15)

    add_heading(doc, "Diagrama de Casos d'Ús", level=2)
    add_body(doc,
        "El sistema identifica quatre actors principals: l'Usuari Anònim (visitant), l'Usuari Registrat "
        "(pla bàsic), l'Usuari Premium (amb accés a entrenador) i l'Administrador. A continuació "
        "es representen els principals casos d'ús:"
    )
    add_screenshot_placeholder(doc, "[ DIAGRAMA DE CASOS D'ÚS: Actors i Funcionalitats Principals ]", width_cm=15)

    add_heading(doc, "Diagrama Entitat-Relació (E-R)", level=2)
    add_body(doc,
        "La base de dades conté les entitats principals i les seves relacions. A continuació es mostra "
        "el diagrama entitat-relació del sistema:"
    )
    add_screenshot_placeholder(doc, "[ DIAGRAMA E-R: Taules i Relacions de la Base de Dades ]", width_cm=15)
    doc.add_page_break()

    # ─── 5. TECNOLOGIES ──────────────────────────────────────────────
    add_heading(doc, "5. Estudi de Tecnologies i Eines", level=1)
    add_body(doc,
        "Abans de iniciar la implementació, vam analitzar les alternatives disponibles per a cada "
        "capa de l'aplicació. La tria va estar guiada per criteris tècnics com la maduresa de la "
        "tecnologia, la seva adopció en l'industria, la corba d'aprenentatge i la compatibilitat "
        "entre components."
    )

    add_heading(doc, "5.1 Back-end – Node.js + Express", level=2)
    add_body(doc,
        "Per al servidor, es van avaluar tres opcions principals: PHP/Laravel, Python/Django i "
        "Node.js/Express. La tria final va recaure en Node.js per les següents raons:"
    )
    for b in [
        "JavaScript al back-end i front-end redueix el canvi de context i unifica l'ecosistema.",
        "Express és minimalista i flexible, ideal per construir una API RESTful a mida.",
        "L'ecosistema npm ofereix una gran varietat de paquets madurs (passport, multer, nodemailer).",
        "Node.js té molt bona adopció professional, la qual cosa facilita la recerca de documentació.",
        "El model no-bloquejant és eficient per gestionar múltiples peticions concurrents.",
    ]:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_info_table(doc, [
        ["Express 5.1.0", "Framework web minimalista per a Node.js"],
        ["mysql2 3.15.3", "Driver MySQL amb suport promises i pool de connexions"],
        ["jsonwebtoken 9.0.2", "Generació i verificació de tokens JWT"],
        ["passport.js", "Middleware d'autenticació (estratègia Google OAuth 2.0)"],
        ["bcryptjs", "Hash i verificació de contrasenyes"],
        ["multer", "Gestió d'uploads de fitxers (fotos de perfil)"],
        ["nodemailer", "Enviament d'emails (formulari de contacte)"],
        ["helmet", "Capçaleres de seguretat HTTP"],
        ["express-rate-limit", "Limitació de peticions per IP"],
        ["swagger-ui-express", "Documentació automàtica de l'API"],
        ["dotenv", "Gestió de variables d'entorn"],
    ], headers=["Paquet", "Funció"])

    add_heading(doc, "5.2 Front-end – React 19 + Vite + Tailwind CSS", level=2)
    add_body(doc,
        "Per al client web, es van comparar Vue 3, Angular i React. La decisió va ser React per "
        "la seva flexibilitat, la gran comunitat, la compatibilitat amb Three.js per als elements 3D "
        "i per la robustesa de React Router per a la navegació SPA."
    )
    add_info_table(doc, [
        ["React 19.2.0", "Biblioteca per a interfícies d'usuari basada en components"],
        ["React Router DOM 7", "Enrutament declaratiu per a SPA"],
        ["Vite 7.3.1", "Bundler ultraràpid amb HMR (Hot Module Replacement)"],
        ["Tailwind CSS 4.1.18", "Framework CSS utility-first per a disseny responsiu"],
        ["Axios 1.13.5", "Client HTTP per a peticions a l'API"],
        ["Framer Motion 12.35", "Biblioteca d'animacions declaratives"],
        ["React-hot-toast 2.6", "Notificacions tipus toast"],
        ["Recharts 3.8.1", "Gràfics de progrés i estadístiques"],
        ["Three.js + R3F", "Renderitzat 3D per a models anatòmics"],
    ], headers=["Paquet", "Funció"])

    add_heading(doc, "5.3 Base de Dades – MySQL 8.0", level=2)
    add_body(doc,
        "Es va triar MySQL per la seva fiabilitat, el suport a transaccions ACID, la compatibilitat "
        "amb hosting compartit i la familiaritat de l'equip. Es va descartar PostgreSQL per no "
        "aportar avantatges decisius per al nostre cas d'ús, i MongoDB per la naturalesa relacional "
        "de les dades (usuaris, recetes, exercicis, rols)."
    )

    add_heading(doc, "5.4 Infraestructura – Docker + Docker Compose", level=2)
    add_body(doc,
        "L'ús de Docker permet garantir que l'entorn de desenvolupament sigui idèntic al de "
        "producció. Docker Compose orquestra els dos serveis principals: el servidor Node.js "
        "i la base de dades MySQL."
    )
    add_code_block(doc,
        "# docker-compose.yml (esquema)\n"
        "services:\n"
        "  db:\n"
        "    image: mysql:8.0\n"
        "    environment:\n"
        "      MYSQL_ROOT_PASSWORD: secret\n"
        "      MYSQL_DATABASE: fitmeal\n"
        "  backend:\n"
        "    build: .\n"
        "    depends_on: [db]\n"
        "    ports: ['3000:3000']\n"
        "    env_file: .env"
    )
    doc.add_page_break()

    # ─── 6. IMPLEMENTACIÓ ─────────────────────────────────────────────
    add_heading(doc, "6. Implementació", level=1)
    add_body(doc,
        "La implementació s'ha estructurat en vuit processos que cobreixen les funcionalitats "
        "principals de l'aplicació. Per a cadascun s'especifica l'objectiu, les possibles solucions "
        "estudiades, els riscos identificats, la implementació realitzada i les proves unitàries "
        "dutes a terme."
    )

    # ─── PROCÉS 1 – AUTENTICACIÓ ──────────────────────────────────────
    add_heading(doc, "6.1 Procés 1 – Autenticació i Gestió de Rols", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Implementar un sistema d'autenticació segur que permeti el registre i inici de sessió "
        "amb email/contrasenya i via OAuth 2.0 de Google. Gestionar quatre rols d'usuari (Admin, "
        "Normal, Premium, Entrenador) amb protecció de rutes tant al front-end com al back-end."
    )

    add_heading(doc, "b) Estudi de Solucions", level=3)
    add_info_table(doc, [
        ["Sessions + Cookies", "Senzill, però dificulta l'escalabilitat horitzontal i les APIs stateless."],
        ["JWT + localStorage", "Fàcil implementació, però vulnerable a XSS."],
        ["JWT + httpOnly Cookie", "✅ Triat. Combina la simplicitat de JWT amb la seguretat de les httpOnly cookies."],
        ["OAuth únicament", "No permet registre propi. Dependència de tercers."],
    ], headers=["Solució", "Avaluació"])

    add_heading(doc, "c) Riscos", level=3)
    for r in [
        "Robatori de tokens si s'usen cookies SameSite mal configurades.",
        "Atacs de força bruta sobre el formulari de login.",
        "Tokens caducats que no es redirigeixen correctament.",
        "Configuració incorrecta del callback OAuth en producció.",
    ]:
        add_bullet(doc, r)

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "El sistema utilitza JWT amb una expiració de 24 hores. Les credencials de l'usuari "
        "es validen al back-end i el token es retorna en una httpOnly cookie (no accessible "
        "des de JavaScript). El middleware auth.js verifica el token en cada petició protegida "
        "i comprova el rol necessari per accedir al recurs."
    )
    add_code_block(doc,
        "// middleware/auth.js – verifyToken\n"
        "const verifyToken = (req, res, next) => {\n"
        "  const token = req.cookies.token || req.headers.authorization?.split(' ')[1];\n"
        "  if (!token) return res.status(401).json({ error: 'No token provided' });\n"
        "  try {\n"
        "    req.user = jwt.verify(token, process.env.JWT_SECRET);\n"
        "    next();\n"
        "  } catch (err) {\n"
        "    return res.status(401).json({ error: 'Invalid token' });\n"
        "  }\n"
        "};\n\n"
        "// Protecció per rol\n"
        "const requireRole = (...roles) => (req, res, next) => {\n"
        "  if (!roles.includes(req.user.id_rol)) {\n"
        "    return res.status(403).json({ error: 'Insufficient permissions' });\n"
        "  }\n"
        "  next();\n"
        "};"
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Login i Pàgina de Registre ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Flux OAuth Google – Pantalla de Selecció de Compte ]")

    add_heading(doc, "e) Proves Unitàries", level=3)
    add_info_table(doc, [
        ["T-AUTH-01", "Registre amb email vàlid", "Retorna 201 + JWT cookie", "✅ Passa"],
        ["T-AUTH-02", "Login amb credencials incorrectes", "Retorna 401", "✅ Passa"],
        ["T-AUTH-03", "Accés a ruta protegida sense token", "Retorna 401", "✅ Passa"],
        ["T-AUTH-04", "Accés a ruta d'admin amb rol=2", "Retorna 403", "✅ Passa"],
        ["T-AUTH-05", "Rate limit: >10 intents en 15min", "Retorna 429", "✅ Passa"],
        ["T-AUTH-06", "Login amb Google OAuth", "Redirigeix correctament", "✅ Passa"],
    ], headers=["Test", "Descripció", "Resultat Esperat", "Estat"])
    doc.add_page_break()

    # ─── PROCÉS 2 – RECETES ──────────────────────────────────────────
    add_heading(doc, "6.2 Procés 2 – Gestió de Recetes i Nutrició", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Proporcionar un catàleg de recetes amb informació nutricional completa (calories, proteïnes, "
        "carbohidrats, grasses, temps de preparació). Permetre als usuaris cercar, filtrar i guardar "
        "recetes als seus favorits. L'administrador ha de poder gestionar el catàleg complet."
    )

    add_heading(doc, "b) Estudi de Solucions", level=3)
    add_body(doc,
        "Es va valorar l'ús d'una API externa de nutrició com Edamam o USDA FoodData Central per "
        "obtenir dades nutricionals automàticament. No obstant, es va decidir mantenir una base de "
        "dades pròpia per tenir control total sobre el contingut i evitar dependències externes "
        "que podrien afectar la disponibilitat o el cost del servei."
    )

    add_heading(doc, "c) Riscos", level=3)
    for r in [
        "Dades nutricionals inexactes si s'introdueixen manualment.",
        "Càrrega excessiva de la pàgina si es recuperen totes les recetes sense paginació.",
        "Imatges d'alta resolució que alenteixen la càrrega inicial.",
    ]:
        add_bullet(doc, r)

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "La taula recetes conté: id_receta, titulo, calorias, proteina, carbohidratos, grasas, "
        "tiempo, tipo, imagen, ingredientes i instrucciones. L'API exposa endpoints CRUD protegits "
        "per rol d'administrador. El front-end implementa filtratge per tipus (breakfast, lunch, "
        "dinner, snack) i cercador per títol."
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Llistat de Recetes amb Filtres ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Detall de Recepta (Ingredients + Valors Nutricionals) ]")

    add_heading(doc, "e) Proves", level=3)
    add_info_table(doc, [
        ["T-REC-01", "GET /api/recipes", "Retorna array de recetes", "✅"],
        ["T-REC-02", "GET /api/recipes/:id", "Retorna recepta per ID", "✅"],
        ["T-REC-03", "POST /api/recipes sense token admin", "Retorna 403", "✅"],
        ["T-REC-04", "Filtrat per tipus 'breakfast'", "Llista filtrada correctament", "✅"],
        ["T-REC-05", "Afegir/eliminar favorit", "BD actualitzada correctament", "✅"],
    ], headers=["Test", "Acció", "Resultat Esperat", "Estat"])
    doc.add_page_break()

    # ─── PROCÉS 3 – EXERCICIS ────────────────────────────────────────
    add_heading(doc, "6.3 Procés 3 – Gestió d'Exercicis i Entrenament", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Oferir un catàleg d'exercicis organitzats per grup muscular amb nivell de dificultat, "
        "descripció i punts clau. Permetre el seguiment del progrés (sèries, repeticions, pes) "
        "i la visualització de models anatòmics 3D per identificar els grups musculars."
    )

    add_heading(doc, "b) Estudi de Solucions", level=3)
    add_body(doc,
        "Per a la visualització anatòmica es va valorar l'ús d'imatges estàtiques vs. un model 3D "
        "interactiu. Es va triar Three.js + React Three Fiber per oferir una experiència visual "
        "diferencial i educativa que permet a l'usuari identificar exactament quins músculs treballa "
        "cada exercici."
    )

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "La taula ejercicios conté: id, musculo_id, titulo, dificultad, imagen, descripcion, tipo "
        "i puntos_clave. L'endpoint GET /api/exercises/:muscleName filtra per grup muscular. "
        "El sistema de progrés emmagatzema date, series, repeticiones i peso per usuari i exercici."
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina d'Exercicis – Selecció per Grup Muscular ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Model 3D Anatòmic + Detall d'Exercici ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Formulari de Registre de Progrés ]")

    add_heading(doc, "e) Proves", level=3)
    add_info_table(doc, [
        ["T-EX-01", "GET /api/exercises", "Llista completa d'exercicis", "✅"],
        ["T-EX-02", "GET /api/exercises/pectorals", "Filtra per múscul", "✅"],
        ["T-EX-03", "POST /api/progress-exercises", "Registra progrés correctament", "✅"],
        ["T-EX-04", "GET /api/favorites-exercises", "Llista favorits de l'usuari", "✅"],
    ], headers=["Test", "Acció", "Resultat Esperat", "Estat"])
    doc.add_page_break()

    # ─── PROCÉS 4 – PERFIL ───────────────────────────────────────────
    add_heading(doc, "6.4 Procés 4 – Perfil d'Usuari i Onboarding", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Recollir les dades inicials de l'usuari (pes, alçada, gènere, objectiu, nivell d'activitat, "
        "preferències alimentàries) a través d'un onboarding guiat per calcular automàticament "
        "les macros recomanades. Permetre l'actualització del perfil i la pujada de foto."
    )

    add_heading(doc, "b) Estudi de Solucions", level=3)
    add_body(doc,
        "Es va valorar demanar les dades en el moment del registre o en un onboarding posterior. "
        "Es va triar el flux d'onboarding post-registre per no allargar el formulari inicial i "
        "reduir l'abandonament durant el registre. L'onboarding es marca com a complet (onboarding_completado=1) "
        "i no es torna a mostrar."
    )

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "La calculadora de macros utilitza la fórmula de Mifflin-St Jeor per calcular el TMB "
        "(Taxa Metabòlica Basal) i aplica el factor d'activitat per obtenir les calories totals "
        "diàries. Distribueix les macros segons l'objectiu: pèrdua de pes, manteniment o guany muscular."
    )
    add_code_block(doc,
        "// utils/macrosCalculator.js\n"
        "export function calculateMacros(weight, height, age, gender, activityLevel, goal) {\n"
        "  // Mifflin-St Jeor BMR\n"
        "  const bmr = gender === 'male'\n"
        "    ? 10 * weight + 6.25 * height - 5 * age + 5\n"
        "    : 10 * weight + 6.25 * height - 5 * age - 161;\n"
        "  const activityFactors = { sedentary:1.2, light:1.375, moderate:1.55, active:1.725 };\n"
        "  const tdee = bmr * activityFactors[activityLevel];\n"
        "  // Goal adjustments\n"
        "  const calories = goal === 'lose' ? tdee - 500 : goal === 'gain' ? tdee + 300 : tdee;\n"
        "  return { calories, protein: weight * 2, carbs: ..., fat: ... };\n"
        "}"
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Formulari d'Onboarding Pas a Pas ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Perfil amb Dades i Foto ]")
    doc.add_page_break()

    # ─── PROCÉS 5 – BOTIGA ───────────────────────────────────────────
    add_heading(doc, "6.5 Procés 5 – Botiga i Carret de la Compra", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Implementar una botiga de productes de fitness (suplements, roba, equipament) amb carret "
        "de la compra, gestió de quantitats i talla, i procés de checkout. La persistència del "
        "carret es gestiona localment per simplificar la implementació."
    )

    add_heading(doc, "b) Estudi de Solucions", level=3)
    add_info_table(doc, [
        ["Carret al backend (BD)", "Persistència entre dispositius, però requereix més endpoints i complexitat."],
        ["Carret a localStorage", "✅ Triat. Senzill, ràpid, sense peticions al servidor. Suficient per MVP."],
        ["Carret a Redux/Zustand", "Interessant per apps grans, però innecessari per a la nostra escala."],
    ], headers=["Solució", "Avaluació"])

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "CartContext.jsx gestiona l'estat global del carret mitjançant useReducer i localStorage. "
        "Cada producte s'identifica per id + talla per permetre la mateixa referència en múltiples "
        "talles. El checkout mostra un formulari d'adreça de lliurament (simulat, sense passerella "
        "de pagament real)."
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Productes de la Botiga ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Detall de Producte ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Carret de la Compra i Checkout ]")
    doc.add_page_break()

    # ─── PROCÉS 6 – ADMIN ────────────────────────────────────────────
    add_heading(doc, "6.6 Procés 6 – Dashboard d'Administrador", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Proporcionar una interfície d'administració per gestionar tots els recursos del sistema: "
        "usuaris, recetes i exercicis. Mostrar estadístiques generals de la plataforma."
    )

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "El dashboard d'administrador és accessible exclusivament per a usuaris amb rol=1. "
        "Utilitza AdminProtectedRoute per redirigir als usuaris sense permís. L'API exposa "
        "GET /api/admin/stats que retorna totals de cada entitat."
    )
    add_info_table(doc, [
        ["GET /api/admin/stats", "Estadístiques: total usuaris, recetes, exercicis, comandes"],
        ["GET /api/admin/users", "Llistat paginat d'usuaris amb filtres"],
        ["GET /api/admin/recipes", "Llistat de recetes amb opcions CRUD"],
        ["GET /api/admin/exercises", "Llistat d'exercicis amb opcions CRUD"],
    ], headers=["Endpoint", "Funció"])
    add_screenshot_placeholder(doc, "[ CAPTURA: Dashboard d'Admin – Estadístiques Globals ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Admin – Gestió d'Usuaris (Llistat + Accions) ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Admin – Gestió de Recetes i Exercicis ]")
    doc.add_page_break()

    # ─── PROCÉS 7 – ENTRENADOR ───────────────────────────────────────
    add_heading(doc, "6.7 Procés 7 – Dashboard d'Entrenador", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Permetre als usuaris amb rol d'entrenador (rol=4) gestionar els seus clients: veure "
        "la llista de clients assignats, crear i assignar rutines d'exercicis personalitzades, "
        "registrar l'entrenament realitzat i fer el seguiment del progrés diari."
    )

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "La taula entrenador_cliente relaciona entrenadors i clients amb estat (activo/inactivo) "
        "i data d'assignació. L'entrenador pot veure la rutina de cada client i afegir exercicis "
        "des del catàleg. L'endpoint POST /api/trainers/log-workout registra el treball realitzat "
        "per data."
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Dashboard d'Entrenador – Llista de Clients ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Vista de Rutina d'un Client + Assignació d'Exercicis ]")
    doc.add_page_break()

    # ─── PROCÉS 8 – CONTACTE ─────────────────────────────────────────
    add_heading(doc, "6.8 Procés 8 – Sistema de Contacte per Email", level=2)

    add_heading(doc, "a) Objectiu", level=3)
    add_body(doc,
        "Implementar un formulari de contacte que enviï un email real a l'equip de FitMeal "
        "i una confirmació automàtica a l'usuari. El sistema no requereix autenticació per "
        "facilitar el contacte de nous usuaris potencials."
    )

    add_heading(doc, "d) Implementació", level=3)
    add_body(doc,
        "S'utilitza Nodemailer configurat amb Gmail SMTP. Quan un usuari envia el formulari, "
        "el back-end genera dos correus: un per l'equip intern amb les dades del contacte, "
        "i un altre de confirmació per a l'usuari que ha enviat el missatge."
    )
    add_code_block(doc,
        "// Exemple configuració Nodemailer\n"
        "const transporter = nodemailer.createTransport({\n"
        "  service: 'gmail',\n"
        "  auth: {\n"
        "    user: process.env.GMAIL_USER,\n"
        "    pass: process.env.GMAIL_APP_PASSWORD  // App password, no la contrasenya real\n"
        "  }\n"
        "});\n\n"
        "// POST /api/contact\n"
        "await transporter.sendMail({\n"
        "  from: process.env.GMAIL_USER,\n"
        "  to: 'fitmeal@gmail.com',\n"
        "  subject: `Nou contacte de ${nombre}`,\n"
        "  html: `<p>${mensaje}</p>`\n"
        "});"
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Contacte amb Formulari ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Email de Confirmació Rebut per l'Usuari ]")
    doc.add_page_break()

    # ─── SCHEMA BD ───────────────────────────────────────────────────
    add_heading(doc, "Estructura de la Base de Dades", level=1)
    add_body(doc,
        "A continuació es detalla l'estructura de les taules principals de la base de dades MySQL. "
        "El sistema utilitza un pool de connexions amb un màxim de 10 connexions concurrents."
    )
    add_screenshot_placeholder(doc, "[ CAPTURA: Esquema SQL – Vista del Workbench o phpMyAdmin ]")

    tables_info = [
        ("usuarios", [
            ("id_usuario", "INT PK AUTO_INCREMENT"),
            ("email", "VARCHAR(100) UNIQUE NOT NULL"),
            ("password_hash", "VARCHAR(255)"),
            ("nombre", "VARCHAR(100)"),
            ("apellidos", "VARCHAR(100)"),
            ("id_rol", "INT FK → roles(id_rol)"),
            ("plan", "ENUM('basico','avanzado','premium')"),
            ("peso", "DECIMAL(5,2)"),
            ("altura", "INT"),
            ("genero", "ENUM('male','female','other')"),
            ("nivel_actividad", "VARCHAR(50)"),
            ("objetivo", "VARCHAR(100)"),
            ("onboarding_completado", "TINYINT(1) DEFAULT 0"),
            ("foto_url", "VARCHAR(255)"),
        ]),
        ("recetas", [
            ("id_receta", "INT PK AUTO_INCREMENT"),
            ("titulo", "VARCHAR(200) NOT NULL"),
            ("calorias", "INT"),
            ("proteina", "DECIMAL(6,2)"),
            ("carbohidratos", "DECIMAL(6,2)"),
            ("grasas", "DECIMAL(6,2)"),
            ("tiempo", "INT (minuts)"),
            ("tipo", "VARCHAR(50)"),
            ("imagen", "VARCHAR(255)"),
            ("ingredientes", "TEXT (JSON)"),
            ("instrucciones", "TEXT"),
        ]),
        ("ejercicios", [
            ("id", "INT PK AUTO_INCREMENT"),
            ("musculo_id", "INT FK → musculos"),
            ("titulo", "VARCHAR(200)"),
            ("dificultad", "ENUM('Baja','Media','Alta')"),
            ("imagen", "VARCHAR(255)"),
            ("descripcion", "TEXT"),
            ("tipo", "VARCHAR(50)"),
            ("puntos_clave", "TEXT (JSON)"),
        ]),
        ("productos", [
            ("id_producto", "INT PK AUTO_INCREMENT"),
            ("nombre_producto", "VARCHAR(200)"),
            ("descripcion", "TEXT"),
            ("precio", "DECIMAL(10,2)"),
            ("stock", "INT"),
            ("id_categoria", "INT FK → categorias_productos"),
            ("imagen_url", "VARCHAR(255)"),
            ("estado", "ENUM('activo','inactivo')"),
        ]),
        ("entrenador_cliente", [
            ("id_entrenador", "INT FK → usuarios"),
            ("id_cliente", "INT FK → usuarios"),
            ("estado", "ENUM('activo','inactivo')"),
            ("fecha_asignacion", "DATETIME"),
        ]),
        ("progreso_ejercicios", [
            ("id", "INT PK AUTO_INCREMENT"),
            ("id_usuario", "INT FK → usuarios"),
            ("id_ejercicio", "INT FK → ejercicios"),
            ("fecha", "DATE"),
            ("series", "INT"),
            ("repeticiones", "INT"),
            ("peso", "DECIMAL(5,2)"),
        ]),
    ]

    for table_name, columns in tables_info:
        add_heading(doc, f"Taula: {table_name}", level=3)
        add_info_table(doc, columns, headers=["Camp", "Tipus / Descripció"])

    doc.add_page_break()

    # ─── FLUXOS ──────────────────────────────────────────────────────
    add_heading(doc, "Diagrames de Flux dels Processos Principals", level=1)

    add_heading(doc, "Flux d'Autenticació", level=2)
    add_code_block(doc,
        "Usuari omple formulari de login\n"
        "  → POST /auth/login\n"
        "    → authController.login() verifica email+password (bcrypt.compare)\n"
        "      → Si OK: generateToken() → JWT a httpOnly cookie\n"
        "        → AuthContext.login() guarda user + token a context\n"
        "          → React Router redirigeix a /perfil o /onboarding\n"
        "      → Si KO: res.status(401) → toast d'error al formulari\n\n"
        "Google OAuth:\n"
        "Usuari clica 'Login amb Google'\n"
        "  → GET /auth/google → Passport redirigeix a Google\n"
        "    → Google autentica i torna al callback\n"
        "      → GET /auth/google/callback\n"
        "        → Passport crea/troba usuari a BD\n"
        "          → JWT cookie → redirigeix al frontend"
    )
    add_screenshot_placeholder(doc, "[ DIAGRAMA: Flux d'Autenticació JWT + OAuth ]")

    add_heading(doc, "Flux de Creació d'una Rutina d'Entrenament", level=2)
    add_code_block(doc,
        "Entrenador accedeix a /entrenador\n"
        "  → GET /api/trainers/clients → Llista clients\n"
        "    → Selecciona client → GET /api/trainers/clients/:id/routine\n"
        "      → Veu rutina actual del client\n"
        "        → Selecciona exercici del catàleg\n"
        "          → POST /api/trainers/assign-exercise\n"
        "            → DB: INSERT a rutina_ejercicios\n"
        "              → Client veu rutina actualitzada a /rutina"
    )
    add_screenshot_placeholder(doc, "[ DIAGRAMA: Flux de Gestió de Rutina Entrenador-Client ]")

    add_heading(doc, "Flux de Compra a la Botiga", level=2)
    add_code_block(doc,
        "Usuari visita /products\n"
        "  → GET /api/products → Llistat de productes\n"
        "    → Clica producte → /products/:id\n"
        "      → Selecciona talla + quantitat\n"
        "        → addToCart() → CartContext + localStorage\n"
        "          → Icona carret actualitza badge\n"
        "            → Usuari va a /cart → Revisa items\n"
        "              → Continua a /checkout\n"
        "                → Omple adreça → Confirma (simulat)"
    )
    add_screenshot_placeholder(doc, "[ DIAGRAMA: Flux de Compra a la Botiga ]")
    doc.add_page_break()

    # ─── 7. MILLORES ─────────────────────────────────────────────────
    add_heading(doc, "7. Millores i Línies de Futur", level=1)

    add_heading(doc, "Si Comencéssim de Nou...", level=2)
    millores_inici = [
        ("Passerella de pagament real", "Integraríem Stripe o PayPal des del primer dia per al checkout, evitant la solució simulada actual."),
        ("TypeScript al back-end", "Hauríem usat TypeScript des d'inici per evitar errors de tipus i millorar el manteniment."),
        ("ORM (Prisma / Sequelize)", "En lloc de SQL raw, un ORM hauria simplificat les consultes i les migracions."),
        ("Testing des del dia 1", "Hauríem implementat tests unitaris i d'integració amb Jest des de l'inici, no al final."),
        ("Disseny mòbil primer", "Algunes pàgines es van dissenyar primer per a escriptori; el mòbil-first hauria evitat retalls posteriors."),
        ("Variables d'entorn unificades", "Separar millor els entorns dev/staging/prod amb fitxers .env específics."),
    ]
    add_info_table(doc, millores_inici, headers=["Àrea", "Millora Proposada"])

    add_heading(doc, "Si Seguíssim Treballant...", level=2)
    millores_futur = [
        ("Passerella de pagament", "Integrar Stripe per a pagaments reals a la botiga i subscripcions premium."),
        ("App mòbil", "Desenvolupar una aplicació nativa amb React Native reutilitzant la lògica de negoci existent."),
        ("IA nutricional", "Integrar un model de llenguatge per generar plans de dieta personalitzats automàticament."),
        ("Tracking avançat", "Gràfics de progrés temporal del pes, composició corporal i records personals."),
        ("Notificacions push", "Sistema de recordatoris per a sessions d'entrenament i ingesta hídrica."),
        ("Social features", "Perfils públics, reptes entre usuaris, sistema de punts i badges."),
        ("Integració amb wearables", "Connexió amb Google Fit, Apple Health i dispositius com Garmin o Fitbit."),
        ("Videotrucades entrenador", "Integrar WebRTC per a sessions en viu entre entrenador i client."),
        ("CI/CD complet", "Pipelines automatitzats amb GitHub Actions, tests automàtics i desplegament continu."),
    ]
    add_info_table(doc, millores_futur, headers=["Funcionalitat", "Descripció"])
    doc.add_page_break()

    # ─── 8. CONCLUSIONS ──────────────────────────────────────────────
    add_heading(doc, "8. Conclusions", level=1)
    add_body(doc,
        "El projecte FitMeal ha suposat un repte tècnic i personal de gran envergadura que ha "
        "permès consolidar i posar en pràctica de forma integrada tots els coneixements adquirits "
        "durant el cicle formatiu de Desenvolupament d'Aplicacions Web."
    )
    add_body(doc,
        "Des del punt de vista tècnic, el projecte ha cobert un espectrum molt ampli de competències: "
        "disseny de bases de dades relacionals, implementació d'una API RESTful amb autenticació "
        "JWT i OAuth, gestió d'arxius, enviament de correus electrònics reals, contenidorització "
        "amb Docker, i desenvolupament d'una SPA completa amb React que inclou animacions, "
        "gràfics, models 3D i un sistema de rols."
    )
    add_body(doc,
        "Un dels aprenentatges més valuosos ha estat la gestió de la complexitat creixent d'un "
        "projecte real. A diferència dels exercicis acadèmics, FitMeal ha requerit decisions "
        "d'arquitectura que tinguessin en compte el manteniment futur, la seguretat i l'experiència "
        "d'usuari de forma simultània. Hem après que no existeix una solució perfecta: sempre hi "
        "ha compromisos entre simplicitat, rendiment i funcionalitat."
    )
    add_body(doc,
        "El treball en equip ha estat fonamental. La coordinació a través de Git (gestió de "
        "branques, resolució de conflictes, pull requests) ha millorat substancialment les nostres "
        "habilitats de col·laboració professional. Hem après a comunicar decisions tècniques, "
        "dividir tasques de forma efectiva i resoldre bloquejos de forma conjunta."
    )
    add_body(doc,
        "Respecte a les limitacions actuals, som conscients que el checkout és simulat, que els "
        "tests automatitzats haurien d'estar millor coberts i que alguns aspectes del disseny mòbil "
        "es podrien polir. Tanmateix, estem satisfets d'haver construït una aplicació funcional, "
        "segura i amb una base sòlida sobre la qual continuar construint."
    )
    add_body(doc,
        "En resum, FitMeal és molt més que un projecte de fi de cicle: és una demostració pràctica "
        "de la nostra capacitat per afrontar el desenvolupament d'una aplicació web real, des de "
        "l'especificació fins al desplegament, amb les eines i metodologies que utilitza la "
        "indústria avui en dia."
    )
    doc.add_page_break()

    # ─── 9. BIBLIOGRAFIA ─────────────────────────────────────────────
    add_heading(doc, "9. Bibliografia / Webgrafia", level=1)
    refs = [
        ("[1]", "MDN Web Docs – JavaScript Reference", "https://developer.mozilla.org"),
        ("[2]", "React Documentation", "https://react.dev"),
        ("[3]", "Express.js Documentation", "https://expressjs.com"),
        ("[4]", "Node.js Documentation", "https://nodejs.org/en/docs"),
        ("[5]", "MySQL 8.0 Reference Manual", "https://dev.mysql.com/doc/refman/8.0/en/"),
        ("[6]", "Tailwind CSS Documentation", "https://tailwindcss.com/docs"),
        ("[7]", "Passport.js Documentation", "https://www.passportjs.org/docs/"),
        ("[8]", "JSON Web Tokens (JWT) – RFC 7519", "https://datatracker.ietf.org/doc/html/rfc7519"),
        ("[9]", "Docker Documentation", "https://docs.docker.com"),
        ("[10]", "Vite Documentation", "https://vitejs.dev/guide/"),
        ("[11]", "React Three Fiber Documentation", "https://docs.pmnd.rs/react-three-fiber"),
        ("[12]", "Framer Motion Documentation", "https://www.framer.com/motion/"),
        ("[13]", "OWASP Top 10 – Security Risks", "https://owasp.org/www-project-top-ten/"),
        ("[14]", "Mifflin-St Jeor BMR Equation", "American Journal of Clinical Nutrition, 1990"),
        ("[15]", "Nodemailer Documentation", "https://nodemailer.com/about/"),
        ("[16]", "Recharts Documentation", "https://recharts.org/en-US/api"),
        ("[17]", "Swagger / OpenAPI Specification", "https://swagger.io/specification/"),
    ]
    add_info_table(doc, refs, headers=["Ref.", "Recurs", "Enllaç / Referència"])
    doc.add_page_break()

    # ─── 10. ANNEXOS ─────────────────────────────────────────────────
    add_heading(doc, "10. Annexos", level=1)

    add_heading(doc, "Annex A – Estructura de Fitxers del Projecte", level=2)
    add_code_block(doc,
        "FitMealProyecto/\n"
        "├── config/\n"
        "│   ├── database.js          # Pool connexions MySQL\n"
        "│   └── passport.js          # Estratègia OAuth Google\n"
        "├── controllers/\n"
        "│   ├── authController.js    # Login, register, verifyToken\n"
        "│   ├── userController.js    # CRUD usuaris\n"
        "│   ├── recipeController.js  # CRUD recetes\n"
        "│   ├── exerciseController.js\n"
        "│   ├── productController.js\n"
        "│   ├── trainerController.js # Lògica entrenador-client\n"
        "│   └── adminController.js   # Estadístiques admin\n"
        "├── middleware/\n"
        "│   ├── auth.js              # verifyToken, requireRole\n"
        "│   ├── upload.js            # Multer (fotos perfil)\n"
        "│   └── log.js               # Logger de peticions\n"
        "├── models/                  # Funcions de consulta BD\n"
        "├── routes/                  # Definició endpoints\n"
        "├── frontend/\n"
        "│   └── src/\n"
        "│       ├── api/axios.js     # Client HTTP configurat\n"
        "│       ├── context/         # AuthContext, CartContext\n"
        "│       ├── pages/           # Pàgines (Login, Home, Perfil...)\n"
        "│       ├── components/      # Navbar, Footer, ProtectedRoute\n"
        "│       └── utils/           # macrosCalculator.js\n"
        "├── index.js                 # Punt d'entrada servidor\n"
        "├── docker-compose.yml\n"
        "└── .env.example"
    )

    add_heading(doc, "Annex B – Variables d'Entorn Necessàries (.env)", level=2)
    add_code_block(doc,
        "# Base de dades\n"
        "DB_HOST=localhost\n"
        "DB_USER=root\n"
        "DB_PASSWORD=****\n"
        "DB_NAME=fitmeal\n"
        "DB_PORT=3306\n\n"
        "# Servidor\n"
        "PORT=3000\n"
        "NODE_ENV=development\n\n"
        "# JWT i Sessions\n"
        "JWT_SECRET=clau_secreta_molt_llarga\n"
        "SESSION_SECRET=clau_sessio\n\n"
        "# OAuth Google\n"
        "GOOGLE_CLIENT_ID=xxx.apps.googleusercontent.com\n"
        "GOOGLE_CLIENT_SECRET=xxx\n"
        "GOOGLE_CALLBACK_URL=http://localhost:3000/auth/google/callback\n\n"
        "# Email (Nodemailer)\n"
        "GMAIL_USER=fitmeal@gmail.com\n"
        "GMAIL_APP_PASSWORD=xxxx_xxxx_xxxx\n\n"
        "# CORS\n"
        "FRONTEND_URL=http://localhost:5173\n"
        "ALLOWED_ORIGINS=http://localhost:5173,http://localhost:3000"
    )

    add_heading(doc, "Annex C – Endpoints de l'API REST (Resum)", level=2)
    endpoints = [
        ("POST", "/auth/register", "No", "Registrar nou usuari"),
        ("POST", "/auth/login", "No", "Inici de sessió"),
        ("GET", "/auth/verify", "JWT", "Verificar token actiu"),
        ("GET", "/auth/google", "No", "Iniciar OAuth Google"),
        ("GET", "/api/recipes", "No", "Llistat de recetes"),
        ("GET", "/api/recipes/:id", "No", "Detall recepta"),
        ("POST", "/api/recipes", "Admin", "Crear recepta"),
        ("PUT", "/api/recipes/:id", "Admin", "Actualitzar recepta"),
        ("DELETE", "/api/recipes/:id", "Admin", "Eliminar recepta"),
        ("GET", "/api/exercises", "No", "Llistat exercicis"),
        ("GET", "/api/exercises/:muscleName", "No", "Filtrar per múscul"),
        ("GET", "/api/products", "No", "Llistat productes"),
        ("GET", "/api/favorites", "JWT", "Recetes favorites"),
        ("POST", "/api/favorites", "JWT", "Afegir favorita"),
        ("GET", "/api/progress-exercises", "JWT", "Progrés entrenaments"),
        ("POST", "/api/progress-exercises", "JWT", "Registrar progrés"),
        ("GET", "/api/trainers/clients", "Entrenador", "Clients assignats"),
        ("POST", "/api/trainers/assign-exercise", "Entrenador", "Assignar exercici"),
        ("GET", "/api/admin/stats", "Admin", "Estadístiques globals"),
        ("POST", "/api/contact", "No", "Enviar email contacte"),
    ]
    add_info_table(doc, endpoints, headers=["Mètode", "Endpoint", "Auth", "Descripció"])

    add_heading(doc, "Annex D – Captures de Pantalla del Projecte", level=2)
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina Principal (Home) – Hero Section ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina Principal – Secció de Plans de Subscripció ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina Principal – Secció de Funcionalitats ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Login – Versió Escriptori ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Login – Versió Mòbil (Responsiu) ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Dashboard d'Administrador complet ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Pàgina de Perfil d'Usuari completa ]")
    add_screenshot_placeholder(doc, "[ CAPTURA: Logo FitMeal – Variants (blanc, color) ]")

    # ─── SAVE ────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Document generat correctament: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_document()
